"""
Document service for fetching and converting Google Drive documents
"""
import re
from typing import Dict
from googleapiclient.errors import HttpError
from app.services.drive_service import DriveService
import logging
from cachetools import TTLCache

logger = logging.getLogger(__name__)

# Cache for 1 hour (3600 seconds), max 100 documents
document_cache = TTLCache(maxsize=100, ttl=3600)


class DocumentService:
    """Service for fetching Google Drive documents"""
    
    def __init__(self):
        self.service = None
        self._initialize_service()
    
    def _initialize_service(self):
        """Initialize Google Drive API service using existing DriveService"""
        try:
            drive_service = DriveService()
            self.service = drive_service.service
            logger.info("Google Drive service initialized successfully for document viewer")
        except Exception as e:
            logger.error(f"Failed to initialize Google Drive service: {e}")
            logger.warning("Document viewer will not work without valid credentials")
            # Don't raise - allow the app to start
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        import html
        return html.escape(text) if text else ''
    
    def extract_file_id(self, url: str) -> str:
        """
        Extract Google Drive file ID from URL
        
        Supports formats:
        - https://docs.google.com/document/d/{FILE_ID}/edit
        - https://drive.google.com/file/d/{FILE_ID}/view
        - https://docs.google.com/spreadsheets/d/{FILE_ID}/edit
        """
        patterns = [
            r'/d/([a-zA-Z0-9-_]+)',
            r'id=([a-zA-Z0-9-_]+)',
            r'/file/d/([a-zA-Z0-9-_]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        raise ValueError(f"Could not extract file ID from URL: {url}")

    async def get_raw_document(self, file_id: str) -> Dict[str, bytes]:
        """
        Get raw document bytes from Google Drive.
        For Google native formats, exports to PDF.
        For uploaded files, downloads directly.
        
        Returns:
            Dict with keys: content (bytes), mime_type, file_name
        """
        if not self.service:
            raise PermissionError("Google Drive service is not available.")
        
        try:
            # Get file metadata
            file_metadata = self.service.files().get(
                fileId=file_id,
                fields='name,mimeType'
            ).execute()
            
            file_name = file_metadata.get('name', 'Document')
            mime_type = file_metadata.get('mimeType', '')
            
            logger.info(f"Fetching raw document: {file_name} (Type: {mime_type})")
            
            # Google native formats - export to PDF
            if 'google-apps' in mime_type:
                content = self.service.files().export(
                    fileId=file_id,
                    mimeType='application/pdf'
                ).execute()
                return {
                    'content': content,
                    'mime_type': 'application/pdf',
                    'file_name': file_name + '.pdf'
                }
            else:
                # Uploaded files - download directly
                content = self.service.files().get_media(fileId=file_id).execute()
                return {
                    'content': content,
                    'mime_type': mime_type,
                    'file_name': file_name
                }
                
        except HttpError as e:
            if e.resp.status == 404:
                raise ValueError("Document not found")
            elif e.resp.status == 403:
                raise PermissionError("Access denied to document")
            else:
                logger.error(f"Google Drive API error: {e}")
                raise Exception(f"Failed to fetch document: {e}")
    
    async def get_document(self, url: str) -> Dict[str, str]:
        """
        Fetch and convert Google Drive document to HTML
        
        Returns:
            Dict with keys: content, mime_type, file_name
        """
        if not self.service:
            raise PermissionError("Google Drive service is not available. Please configure valid credentials.")
        
        file_id = self.extract_file_id(url)
        
        # Check cache
        if file_id in document_cache:
            logger.info(f"Returning cached document: {file_id}")
            return document_cache[file_id]
        
        try:
            # Get file metadata
            file_metadata = self.service.files().get(
                fileId=file_id,
                fields='name,mimeType'
            ).execute()
            
            file_name = file_metadata.get('name', 'Document')
            mime_type = file_metadata.get('mimeType', '')
            
            logger.info(f"Fetching document: {file_name} (ID: {file_id}, Type: {mime_type})")
            
            # Export based on document type
            if 'google-apps.document' in mime_type:
                # Google Docs → HTML
                content = self.service.files().export(
                    fileId=file_id,
                    mimeType='text/html'
                ).execute().decode('utf-8')
                export_mime = 'text/html'
                
            elif 'google-apps.spreadsheet' in mime_type:
                # Google Sheets → HTML
                content = self.service.files().export(
                    fileId=file_id,
                    mimeType='text/html'
                ).execute().decode('utf-8')
                export_mime = 'text/html'
                
            elif 'google-apps.presentation' in mime_type:
                # Google Slides → HTML (slides as images)
                content = self.service.files().export(
                    fileId=file_id,
                    mimeType='text/html'
                ).execute().decode('utf-8')
                export_mime = 'text/html'
                
            elif 'pdf' in mime_type.lower() or 'application/pdf' in mime_type:
                # PDF → Download and embed as base64 data URL
                import base64
                pdf_content = self.service.files().get_media(fileId=file_id).execute()
                pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
                content = f'''
                <html>
                <head>
                    <style>
                        body {{ margin: 0; padding: 0; height: 100vh; overflow: hidden; }}
                        iframe, embed, object {{ width: 100%; height: 100%; border: none; }}
                    </style>
                </head>
                <body>
                    <embed src="data:application/pdf;base64,{pdf_base64}" type="application/pdf" width="100%" height="100%">
                </body>
                </html>
                '''
                export_mime = 'text/html'
                
            elif any(ms_type in mime_type for ms_type in [
                'application/vnd.openxmlformats-officedocument',
                'application/vnd.ms-',
                'application/msword'
            ]):
                # Microsoft Office files → Download and convert to HTML
                import base64
                import io
                
                # Download the file
                file_content = self.service.files().get_media(fileId=file_id).execute()
                
                if 'wordprocessingml' in mime_type or 'msword' in mime_type:
                    # Word documents - convert to HTML using mammoth (best quality)
                    try:
                        import mammoth
                        result = mammoth.convert_to_html(io.BytesIO(file_content))
                        content = result.value
                        
                        # Log any conversion warnings
                        if result.messages:
                            for msg in result.messages:
                                logger.warning(f"Mammoth conversion: {msg}")
                                
                    except Exception as mammoth_error:
                        logger.warning(f"Mammoth conversion failed: {mammoth_error}, trying python-docx")
                        # Fallback to python-docx
                        try:
                            from docx import Document
                            doc = Document(io.BytesIO(file_content))
                            
                            html_parts = ['<div class="docx-content">']
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    html_parts.append(f'<p>{self._escape_html(para.text)}</p>')
                            
                            for table in doc.tables:
                                html_parts.append('<table class="docx-table">')
                                for row in table.rows:
                                    html_parts.append('<tr>')
                                    for cell in row.cells:
                                        html_parts.append(f'<td>{self._escape_html(cell.text)}</td>')
                                    html_parts.append('</tr>')
                                html_parts.append('</table>')
                            
                            html_parts.append('</div>')
                            content = '\\n'.join(html_parts)
                        except Exception as docx_error:
                            logger.error(f"Failed to parse docx: {docx_error}")
                            file_base64 = base64.b64encode(file_content).decode('utf-8')
                            content = f'''
                            <div style="padding: 20px; text-align: center;">
                                <h2>📄 {file_name}</h2>
                                <p>Click below to download and view this document.</p>
                                <a href="data:{mime_type};base64,{file_base64}" 
                                   download="{file_name}" 
                                   style="display: inline-block; padding: 15px 30px; background: #0047AC; color: white; text-decoration: none; border-radius: 5px;">
                                    ⬇️ Download Document
                                </a>
                            </div>
                            '''
                        
                elif 'spreadsheetml' in mime_type:
                    # Excel files - download option
                    file_base64 = base64.b64encode(file_content).decode('utf-8')
                    content = f'''
                    <div style="padding: 20px; text-align: center;">
                        <h2>📊 {file_name}</h2>
                        <p>Excel spreadsheets are best viewed in Microsoft Excel.</p>
                        <a href="data:{mime_type};base64,{file_base64}" 
                           download="{file_name}" 
                           style="display: inline-block; padding: 15px 30px; background: #0047AC; color: white; text-decoration: none; border-radius: 5px;">
                            ⬇️ Download Spreadsheet
                        </a>
                    </div>
                    '''
                    
                elif 'presentationml' in mime_type:
                    # PowerPoint files - download option
                    file_base64 = base64.b64encode(file_content).decode('utf-8')
                    content = f'''
                    <div style="padding: 20px; text-align: center;">
                        <h2>📊 {file_name}</h2>
                        <p>PowerPoint presentations are best viewed in Microsoft PowerPoint.</p>
                        <a href="data:{mime_type};base64,{file_base64}" 
                           download="{file_name}" 
                           style="display: inline-block; padding: 15px 30px; background: #0047AC; color: white; text-decoration: none; border-radius: 5px;">
                            ⬇️ Download Presentation
                        </a>
                    </div>
                    '''
                else:
                    # Other Office formats - download link
                    file_base64 = base64.b64encode(file_content).decode('utf-8')
                    content = f'''
                    <div style="padding: 20px; text-align: center;">
                        <h2>📄 {file_name}</h2>
                        <p>This document can be downloaded for viewing.</p>
                        <a href="data:{mime_type};base64,{file_base64}" 
                           download="{file_name}" 
                           style="display: inline-block; padding: 15px 30px; background: #0047AC; color: white; text-decoration: none; border-radius: 5px;">
                            ⬇️ Download Document
                        </a>
                    </div>
                    '''
                export_mime = 'text/html'
                
            else:
                # Unsupported format - provide link to open in Google Drive
                content = f'''
                <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; padding: 20px; }}
                        h3 {{ color: #333; }}
                        a {{ color: #0047AC; text-decoration: none; }}
                        a:hover {{ text-decoration: underline; }}
                    </style>
                </head>
                <body>
                    <h3>Unsupported Document Type</h3>
                    <p>This document type ({mime_type}) cannot be previewed.</p>
                    <p><a href="{url}" target="_blank">Click here to open in Google Drive</a></p>
                </body>
                </html>
                '''
                export_mime = 'text/html'
            
            # Add basic styling to HTML content
            styled_content = f'''
            <html>
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        padding: 20px;
                        max-width: 900px;
                        margin: 0 auto;
                        line-height: 1.6;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                    }}
                    table td, table th {{
                        border: 1px solid #ddd;
                        padding: 8px;
                    }}
                    table th {{
                        background-color: #f2f2f2;
                    }}
                </style>
            </head>
            <body>
                {content}
            </body>
            </html>
            '''
            
            result = {
                'content': styled_content,
                'mime_type': export_mime,
                'file_name': file_name
            }
            
            # Cache the result
            document_cache[file_id] = result
            
            return result
            
        except HttpError as e:
            if e.resp.status == 404:
                raise ValueError("Document not found")
            elif e.resp.status == 403:
                raise PermissionError("Access denied to document")
            else:
                logger.error(f"Google Drive API error: {e}")
                raise Exception(f"Failed to fetch document: {e}")
